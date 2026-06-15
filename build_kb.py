"""
知识库构建工具 — 支持 PDF / DOCX / TXT 多格式
=============================================
用法:
  python build_kb.py            # 增量添加（追加到现有库）
  python build_kb.py --rebuild  # 重建（清空旧库后重新构建）

支持将 docs/ 目录下的所有 .pdf、.docx、.txt 文件
自动读取、分块、向量化后存入 Chroma 向量库。
"""

import shutil
import sys
from pathlib import Path

from langchain_chroma import Chroma
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document


DOCS_DIR = Path("docs")
DB_PATH = Path("chroma_db")
MODEL_NAME = "sentence-transformers/all-MiniLM-L6-v2"

# 分块策略
CHUNK_SIZE = 400
CHUNK_OVERLAP = 80


def load_docx(file_path: Path) -> list[Document]:
    """读取 .docx 文件"""
    from docx import Document as DocxDocument
    doc = DocxDocument(str(file_path))
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return [Document(page_content=full_text, metadata={"source": str(file_path), "type": "docx"})]


def load_all_documents() -> list[Document]:
    """扫描 docs/ 目录，加载所有支持的文档"""
    all_docs: list[Document] = []
    loaders = {
        ".pdf": lambda p: PyPDFLoader(str(p)).load(),
        ".txt": lambda p: TextLoader(str(p), encoding="utf-8").load(),
        ".docx": lambda p: load_docx(p),
    }
    for ext, loader_fn in loaders.items():
        for file_path in DOCS_DIR.glob(f"*{ext}"):
            print(f"  📄 读取: {file_path.name}")
            try:
                docs = loader_fn(file_path)
                for d in docs:
                    d.metadata["source"] = file_path.name
                all_docs.extend(docs)
                print(f"     → {len(docs)} 页/段")
            except Exception as e:
                print(f"     ⚠️  读取失败: {e}")
    return all_docs


def main() -> None:
    rebuild = "--rebuild" in sys.argv

    print("=" * 50)
    print("📚 大学选课助手 — 构建工具")
    print("=" * 50)

    # 重建模式：清空旧库
    if rebuild:
        if DB_PATH.exists():
            print("\n🗑️  清空旧向量库（--rebuild 模式）...")
            shutil.rmtree(DB_PATH)
            print("   已清空")

    print("\n🔍 扫描 docs/ 目录...")
    if not DOCS_DIR.exists():
        DOCS_DIR.mkdir(parents=True)
        print("   📁 docs/ 目录已创建，请放入文档后重试")
        return

    documents = load_all_documents()
    if not documents:
        print("   ⚠️  docs/ 目录中没有找到支持的文档")
        return
    print(f"\n📊 共加载 {len(documents)} 个文档片段")

    print("\n✂️  文本分块中...")
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", "。", "；", "，", " ", ""],
    )
    chunks = splitter.split_documents(documents)
    print(f"   → 分成 {len(chunks)} 个块")

    print("\n🧠 加载 embedding 模型...")
    embeddings = HuggingFaceEmbeddings(model_name=MODEL_NAME)

    print("\n💾 写入 Chroma 向量库...")
    Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        persist_directory=str(DB_PATH),
    )
    print(f"   ✅ 知识库创建成功 → {DB_PATH}/")
    print(f"   📦 共存储 {len(chunks)} 个向量")


if __name__ == "__main__":
    main()
